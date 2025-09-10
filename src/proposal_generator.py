"""
Professional Proposal Generator for ROI Calculator
Creates comprehensive sales proposals with customizable templates and multiple formats
"""

import os
import json
import tempfile
from datetime import datetime, timedelta
from pathlib import Path
import logging
from typing import Dict, Any, List, Optional

# Document generation imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.charts.barcharts import VerticalBarChart

# Chart generation
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import seaborn as sns
import numpy as np

logger = logging.getLogger(__name__)

class ProposalGenerator:
    """Professional proposal generator with multiple format support"""
    
    def __init__(self, roi_results: Dict[str, Any], company_config: Optional[Dict] = None):
        """
        Initialize proposal generator
        
        Args:
            roi_results: ROI calculation results
            company_config: Company branding configuration
        """
        self.roi_results = roi_results
        self.company_config = company_config or self._get_default_company_config()
        self.generated_charts = []
        
        # Set up matplotlib for professional charts
        plt.style.use('seaborn-v0_8-whitegrid')
        sns.set_palette("husl")
        
        # Create temp directory for charts
        self.temp_dir = tempfile.mkdtemp(prefix='roi_proposal_')
        
    def _get_default_company_config(self) -> Dict:
        """Get default company configuration"""
        return {
            'name': 'ROI Solutions Inc.',
            'address': '123 Business Ave, Suite 100\nBusiness City, BC 12345',
            'phone': '+1 (555) 123-4567',
            'email': 'contact@roisolutions.com',
            'website': 'www.roisolutions.com',
            'logo_path': None,
            'primary_color': '#2E5BBA',
            'secondary_color': '#8FA4D3',
            'accent_color': '#D4AF37'
        }
    
    def generate_proposal(self, format_type: str = 'all', template_name: str = 'professional') -> Dict[str, str]:
        """
        Generate proposal in specified format(s)
        
        Args:
            format_type: 'pdf', 'docx', 'html', or 'all'
            template_name: Template to use ('professional', 'modern', 'executive')
            
        Returns:
            Dictionary mapping format to file path
        """
        results = {}
        
        try:
            # Generate charts first
            self._generate_charts()
            
            if format_type in ['pdf', 'all']:
                results['pdf'] = self._generate_pdf_proposal(template_name)
                
            if format_type in ['docx', 'all']:
                results['docx'] = self._generate_docx_proposal(template_name)
                
            if format_type in ['html', 'all']:
                results['html'] = self._generate_html_proposal(template_name)
                
            return results
            
        except Exception as e:
            logger.error(f"Error generating proposal: {str(e)}")
            raise
    
    def _generate_charts(self) -> None:
        """Generate charts for proposal"""
        try:
            # ROI Timeline Chart
            self._create_roi_timeline_chart()
            
            # Cost Breakdown Chart
            self._create_cost_breakdown_chart()
            
            # Savings Projection Chart
            self._create_savings_projection_chart()
            
            # Investment vs Returns Chart
            self._create_investment_returns_chart()
            
            # Risk-Return Analysis Chart
            self._create_risk_return_chart()
            
        except Exception as e:
            logger.error(f"Error generating charts: {str(e)}")
    
    def _create_roi_timeline_chart(self) -> str:
        """Create ROI timeline visualization"""
        fig, ax = plt.subplots(figsize=(12, 6))
        
        # Extract projection data
        projections = self.roi_results.get('projections', {})
        years = []
        roi_values = []
        cumulative_savings = []
        
        for year in ['year_1', 'year_2', 'year_3']:
            if year in projections:
                years.append(int(year.split('_')[1]))
                roi_values.append(projections[year].get('roi_percentage', 0) * 100)
                cumulative_savings.append(projections[year].get('cumulative_savings', 0))
        
        if not years:
            # Use default data if projections not available
            years = [1, 2, 3]
            base_roi = self.roi_results.get('roi_metrics', {}).get('first_year_roi', 0.25) * 100
            roi_values = [base_roi, base_roi * 1.15, base_roi * 1.25]
            base_savings = self.roi_results.get('roi_metrics', {}).get('annual_savings', 50000)
            cumulative_savings = [base_savings, base_savings * 2.1, base_savings * 3.3]
        
        # Create dual-axis chart
        ax2 = ax.twinx()
        
        # ROI percentage line
        line1 = ax.plot(years, roi_values, marker='o', linewidth=3, 
                       color=self.company_config['primary_color'], label='ROI %')
        ax.set_ylabel('ROI Percentage (%)', fontsize=12, fontweight='bold')
        ax.set_ylim(0, max(roi_values) * 1.2)
        
        # Cumulative savings bars
        bars = ax2.bar(years, cumulative_savings, alpha=0.7, 
                      color=self.company_config['secondary_color'], 
                      label='Cumulative Savings ($)')
        ax2.set_ylabel('Cumulative Savings ($)', fontsize=12, fontweight='bold')
        
        # Formatting
        ax.set_xlabel('Year', fontsize=12, fontweight='bold')
        ax.set_title('ROI Timeline & Cumulative Savings', fontsize=16, fontweight='bold', pad=20)
        ax.set_xticks(years)
        ax.grid(True, alpha=0.3)
        
        # Format currency labels
        ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        
        # Add value labels on bars
        for bar, value in zip(bars, cumulative_savings):
            height = bar.get_height()
            ax2.text(bar.get_x() + bar.get_width()/2., height + height*0.01,
                    f'${value:,.0f}', ha='center', va='bottom', fontweight='bold')
        
        # Combined legend
        lines1, labels1 = ax.get_legend_handles_labels()
        lines2, labels2 = ax2.get_legend_handles_labels()
        ax.legend(lines1 + lines2, labels1 + labels2, loc='upper left', frameon=True, 
                 fancybox=True, shadow=True)
        
        plt.tight_layout()
        
        chart_path = os.path.join(self.temp_dir, 'roi_timeline.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        self.generated_charts.append(chart_path)
        return chart_path
    
    def _create_cost_breakdown_chart(self) -> str:
        """Create cost breakdown pie chart"""
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Extract cost data
        inputs = self.roi_results.get('inputs', {})
        costs = {
            'Labor Costs': inputs.get('labor_costs', 0) * 12,  # Monthly to annual
            'Shipping Costs': inputs.get('shipping_costs', 0) * 12,
            'Error Costs': inputs.get('error_costs', 0) * 12,
            'Inventory Costs': inputs.get('inventory_costs', 0) * 12
        }
        
        # Filter out zero costs
        costs = {k: v for k, v in costs.items() if v > 0}
        
        if not costs:
            costs = {'Sample Costs': 100000}  # Fallback
        
        # Create pie chart
        colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7']
        wedges, texts, autotexts = ax.pie(costs.values(), labels=costs.keys(), 
                                         autopct='%1.1f%%', startangle=90,
                                         colors=colors[:len(costs)], 
                                         explode=[0.05] * len(costs))
        
        # Enhance appearance
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
            autotext.set_fontsize(11)
        
        for text in texts:
            text.set_fontsize(12)
            text.set_fontweight('bold')
        
        ax.set_title('Annual Cost Breakdown', fontsize=16, fontweight='bold', pad=20)
        
        # Add total cost in center
        total_cost = sum(costs.values())
        ax.text(0, 0, f'Total\n${total_cost:,.0f}', ha='center', va='center',
               fontsize=14, fontweight='bold', 
               bbox=dict(boxstyle='round,pad=0.5', facecolor='white', alpha=0.8))
        
        plt.tight_layout()
        
        chart_path = os.path.join(self.temp_dir, 'cost_breakdown.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        self.generated_charts.append(chart_path)
        return chart_path
    
    def _create_savings_projection_chart(self) -> str:
        """Create savings projection chart"""
        fig, ax = plt.subplots(figsize=(12, 8))
        
        # Generate monthly savings data for 3 years
        monthly_savings = self.roi_results.get('roi_metrics', {}).get('monthly_savings', 5000)
        months = list(range(1, 37))  # 36 months
        
        # Model increasing savings over time (efficiency improvements)
        base_savings = [monthly_savings * (1 + 0.02 * (m-1)/12) for m in months]  # 2% annual growth
        
        # Add some realistic variance
        import random
        random.seed(42)  # For consistent results
        actual_savings = [s * (0.9 + 0.2 * random.random()) for s in base_savings]
        
        # Calculate cumulative savings
        cumulative_actual = np.cumsum(actual_savings)
        cumulative_projected = np.cumsum(base_savings)
        
        # Create the plot
        ax.plot(months, cumulative_projected, '--', linewidth=2, 
               color=self.company_config['primary_color'], label='Projected Savings')
        ax.plot(months, cumulative_actual, linewidth=3, 
               color=self.company_config['accent_color'], label='Expected Range')
        ax.fill_between(months, cumulative_actual * 0.85, cumulative_actual * 1.15, 
                       alpha=0.2, color=self.company_config['accent_color'])
        
        # Formatting
        ax.set_xlabel('Month', fontsize=12, fontweight='bold')
        ax.set_ylabel('Cumulative Savings ($)', fontsize=12, fontweight='bold')
        ax.set_title('36-Month Savings Projection', fontsize=16, fontweight='bold', pad=20)
        
        # Format y-axis as currency
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        
        # Add milestone markers
        milestones = [12, 24, 36]
        for milestone in milestones:
            if milestone <= len(cumulative_actual):
                value = cumulative_actual[milestone-1]
                ax.axvline(x=milestone, color='red', linestyle=':', alpha=0.7)
                ax.annotate(f'Year {milestone//12}\n${value:,.0f}', 
                           xy=(milestone, value), xytext=(10, 10),
                           textcoords='offset points', 
                           bbox=dict(boxstyle='round,pad=0.3', facecolor='yellow', alpha=0.7),
                           arrowprops=dict(arrowstyle='->', connectionstyle='arc3,rad=0'))
        
        ax.grid(True, alpha=0.3)
        ax.legend(loc='upper left', frameon=True, fancybox=True, shadow=True)
        
        plt.tight_layout()
        
        chart_path = os.path.join(self.temp_dir, 'savings_projection.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        self.generated_charts.append(chart_path)
        return chart_path
    
    def _create_investment_returns_chart(self) -> str:
        """Create investment vs returns comparison chart"""
        fig, ax = plt.subplots(figsize=(10, 6))
        
        investment = self.roi_results.get('inputs', {}).get('service_investment', 50000)
        annual_savings = self.roi_results.get('roi_metrics', {}).get('annual_savings', 60000)
        
        categories = ['Initial Investment', 'Year 1 Returns', 'Year 2 Returns', 'Year 3 Returns']
        values = [
            -investment,  # Negative for investment
            annual_savings,
            annual_savings * 1.1,  # 10% growth
            annual_savings * 1.2   # 20% growth
        ]
        
        # Create bar chart with different colors
        colors = ['red' if v < 0 else self.company_config['primary_color'] for v in values]
        bars = ax.bar(categories, values, color=colors, alpha=0.8)
        
        # Add value labels on bars
        for bar, value in zip(bars, values):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2.,
                   height + (abs(height) * 0.02 if height >= 0 else -abs(height) * 0.02),
                   f'${abs(value):,.0f}', ha='center', 
                   va='bottom' if height >= 0 else 'top',
                   fontweight='bold', fontsize=11)
        
        # Add break-even line
        ax.axhline(y=0, color='black', linestyle='-', linewidth=2, alpha=0.8)
        
        # Formatting
        ax.set_ylabel('Amount ($)', fontsize=12, fontweight='bold')
        ax.set_title('Investment vs Returns Analysis', fontsize=16, fontweight='bold', pad=20)
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        
        # Rotate x-axis labels for better readability
        plt.xticks(rotation=45, ha='right')
        
        # Add net return annotation
        net_3_year = sum(values)
        ax.annotate(f'3-Year Net Return: ${net_3_year:,.0f}', 
                   xy=(0.5, 0.95), xycoords='axes fraction',
                   bbox=dict(boxstyle='round,pad=0.5', facecolor='lightgreen', alpha=0.8),
                   fontsize=12, fontweight='bold', ha='center')
        
        ax.grid(True, alpha=0.3, axis='y')
        plt.tight_layout()
        
        chart_path = os.path.join(self.temp_dir, 'investment_returns.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        self.generated_charts.append(chart_path)
        return chart_path
    
    def _create_risk_return_chart(self) -> str:
        """Create risk-return analysis scatter plot"""
        fig, ax = plt.subplots(figsize=(10, 8))
        
        # Generate scenario data
        scenarios = [
            {'name': 'Conservative', 'risk': 15, 'return': 20, 'color': '#28a745'},
            {'name': 'Most Likely', 'risk': 25, 'return': 35, 'color': self.company_config['primary_color']},
            {'name': 'Optimistic', 'risk': 35, 'return': 55, 'color': self.company_config['accent_color']},
            {'name': 'Market Average', 'risk': 30, 'return': 25, 'color': '#6c757d'},
        ]
        
        for scenario in scenarios:
            ax.scatter(scenario['risk'], scenario['return'], 
                      s=300, color=scenario['color'], alpha=0.8,
                      edgecolors='white', linewidth=2)
            ax.annotate(scenario['name'], 
                       (scenario['risk'], scenario['return']),
                       xytext=(5, 5), textcoords='offset points',
                       fontweight='bold', fontsize=11)
        
        # Add quadrant lines
        ax.axhline(y=30, color='gray', linestyle='--', alpha=0.5)
        ax.axvline(x=25, color='gray', linestyle='--', alpha=0.5)
        
        # Add quadrant labels
        ax.text(12, 50, 'Low Risk\nHigh Return', ha='center', va='center', 
               bbox=dict(boxstyle='round,pad=0.3', facecolor='lightgreen', alpha=0.7))
        ax.text(40, 50, 'High Risk\nHigh Return', ha='center', va='center',
               bbox=dict(boxstyle='round,pad=0.3', facecolor='yellow', alpha=0.7))
        ax.text(12, 15, 'Low Risk\nLow Return', ha='center', va='center',
               bbox=dict(boxstyle='round,pad=0.3', facecolor='lightblue', alpha=0.7))
        ax.text(40, 15, 'High Risk\nLow Return', ha='center', va='center',
               bbox=dict(boxstyle='round,pad=0.3', facecolor='lightcoral', alpha=0.7))
        
        # Formatting
        ax.set_xlabel('Risk Level (%)', fontsize=12, fontweight='bold')
        ax.set_ylabel('Expected Return (%)', fontsize=12, fontweight='bold')
        ax.set_title('Risk-Return Analysis', fontsize=16, fontweight='bold', pad=20)
        
        ax.grid(True, alpha=0.3)
        ax.set_xlim(5, 50)
        ax.set_ylim(5, 60)
        
        plt.tight_layout()
        
        chart_path = os.path.join(self.temp_dir, 'risk_return.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        self.generated_charts.append(chart_path)
        return chart_path
    
    def _generate_pdf_proposal(self, template_name: str) -> str:
        """Generate PDF proposal"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name = self.roi_results.get('inputs', {}).get('company_name', 'Client')
        safe_client = ''.join(c for c in client_name if c.isalnum() or c in (' ', '-', '_')).strip()
        
        filename = f'ROI_Proposal_{safe_client}_{timestamp}.pdf'
        filepath = os.path.join(self.temp_dir, filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=letter,
                               rightMargin=0.75*inch, leftMargin=0.75*inch,
                               topMargin=1*inch, bottomMargin=1*inch)
        
        # Build content
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor(self.company_config['primary_color']),
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=12,
            textColor=colors.HexColor(self.company_config['primary_color'])
        )
        
        # Title page
        story.append(Paragraph("ROI ANALYSIS PROPOSAL", title_style))
        story.append(Spacer(1, 0.5*inch))
        
        story.append(Paragraph(f"Prepared for: {client_name}", styles['Normal']))
        story.append(Paragraph(f"Prepared by: {self.company_config['name']}", styles['Normal']))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        
        story.append(PageBreak())
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        exec_summary = self._generate_executive_summary()
        for paragraph in exec_summary.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 12))
        
        # ROI Analysis section
        story.append(Paragraph("ROI Analysis", heading_style))
        
        # Key metrics table
        roi_data = [
            ['Metric', 'Value'],
            ['Initial Investment', f"${self.roi_results.get('inputs', {}).get('service_investment', 0):,.2f}"],
            ['First Year ROI', f"{self.roi_results.get('roi_metrics', {}).get('first_year_roi', 0)*100:.1f}%"],
            ['Annual Savings', f"${self.roi_results.get('roi_metrics', {}).get('annual_savings', 0):,.2f}"],
            ['Payback Period', f"{self.roi_results.get('roi_metrics', {}).get('payback_period_months', 0):.1f} months"],
            ['3-Year NPV', f"${self.roi_results.get('financial_metrics', {}).get('npv', 0):,.2f}"]
        ]
        
        roi_table = Table(roi_data, colWidths=[2.5*inch, 2*inch])
        roi_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(self.company_config['primary_color'])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(roi_table)
        story.append(Spacer(1, 20))
        
        # Add charts
        if self.generated_charts:
            for i, chart_path in enumerate(self.generated_charts):
                if os.path.exists(chart_path):
                    try:
                        # Add chart with caption
                        chart_names = ['ROI Timeline', 'Cost Breakdown', 'Savings Projection', 
                                     'Investment Returns', 'Risk Analysis']
                        chart_name = chart_names[i] if i < len(chart_names) else f'Chart {i+1}'
                        
                        story.append(Paragraph(chart_name, heading_style))
                        img = Image(chart_path, width=6*inch, height=3*inch)
                        story.append(img)
                        story.append(Spacer(1, 20))
                        
                        if i == 1:  # Add page break after second chart
                            story.append(PageBreak())
                    except Exception as e:
                        logger.warning(f"Could not add chart {chart_path}: {e}")
        
        # Implementation Timeline
        story.append(PageBreak())
        story.append(Paragraph("Implementation Timeline", heading_style))
        timeline = self._generate_implementation_timeline()
        for item in timeline:
            story.append(Paragraph(f"• {item}", styles['Normal']))
            story.append(Spacer(1, 6))
        
        # Cost Breakdown
        story.append(Paragraph("Detailed Cost Analysis", heading_style))
        cost_breakdown = self._generate_cost_breakdown()
        for section, content in cost_breakdown.items():
            story.append(Paragraph(section, styles['Heading3']))
            story.append(Paragraph(content, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Next Steps
        story.append(Paragraph("Next Steps", heading_style))
        next_steps = [
            "Review and approve this ROI analysis",
            "Schedule implementation kickoff meeting",
            "Finalize project timeline and milestones",
            "Begin service deployment and integration",
            "Establish monitoring and tracking procedures"
        ]
        for step in next_steps:
            story.append(Paragraph(f"• {step}", styles['Normal']))
            story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        return filepath
    
    def _generate_docx_proposal(self, template_name: str) -> str:
        """Generate Word document proposal"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name = self.roi_results.get('inputs', {}).get('company_name', 'Client')
        safe_client = ''.join(c for c in client_name if c.isalnum() or c in (' ', '-', '_')).strip()
        
        filename = f'ROI_Proposal_{safe_client}_{timestamp}.docx'
        filepath = os.path.join(self.temp_dir, filename)
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"ROI Analysis Proposal - {client_name}"
        doc.core_properties.author = self.company_config['name']
        doc.core_properties.comments = "Professional ROI analysis and business case"
        
        # Title page
        title = doc.add_heading('ROI ANALYSIS PROPOSAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Client info
        p = doc.add_paragraph()
        p.add_run(f'Prepared for: ').bold = True
        p.add_run(client_name)
        
        p = doc.add_paragraph()
        p.add_run(f'Prepared by: ').bold = True
        p.add_run(self.company_config['name'])
        
        p = doc.add_paragraph()
        p.add_run(f'Date: ').bold = True
        p.add_run(datetime.now().strftime('%B %d, %Y'))
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        exec_summary = self._generate_executive_summary()
        for paragraph in exec_summary.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # ROI Analysis
        doc.add_heading('ROI Analysis', level=1)
        
        # Key metrics table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        metrics = [
            ('Initial Investment', f"${self.roi_results.get('inputs', {}).get('service_investment', 0):,.2f}"),
            ('First Year ROI', f"{self.roi_results.get('roi_metrics', {}).get('first_year_roi', 0)*100:.1f}%"),
            ('Annual Savings', f"${self.roi_results.get('roi_metrics', {}).get('annual_savings', 0):,.2f}"),
            ('Payback Period', f"{self.roi_results.get('roi_metrics', {}).get('payback_period_months', 0):.1f} months"),
            ('3-Year NPV', f"${self.roi_results.get('financial_metrics', {}).get('npv', 0):,.2f}")
        ]
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Add charts
        if self.generated_charts:
            doc.add_heading('Visual Analysis', level=1)
            chart_names = ['ROI Timeline', 'Cost Breakdown', 'Savings Projection', 
                         'Investment Returns', 'Risk Analysis']
            
            for i, chart_path in enumerate(self.generated_charts):
                if os.path.exists(chart_path):
                    try:
                        chart_name = chart_names[i] if i < len(chart_names) else f'Chart {i+1}'
                        doc.add_heading(chart_name, level=2)
                        doc.add_picture(chart_path, width=Inches(6))
                        doc.add_paragraph()
                    except Exception as e:
                        logger.warning(f"Could not add chart {chart_path}: {e}")
        
        # Implementation Timeline
        doc.add_heading('Implementation Timeline', level=1)
        timeline = self._generate_implementation_timeline()
        for item in timeline:
            p = doc.add_paragraph(style='List Bullet')
            p.text = item
        
        # Cost Breakdown
        doc.add_heading('Detailed Cost Analysis', level=1)
        cost_breakdown = self._generate_cost_breakdown()
        for section, content in cost_breakdown.items():
            doc.add_heading(section, level=2)
            doc.add_paragraph(content)
        
        # Next Steps
        doc.add_heading('Next Steps', level=1)
        next_steps = [
            "Review and approve this ROI analysis",
            "Schedule implementation kickoff meeting", 
            "Finalize project timeline and milestones",
            "Begin service deployment and integration",
            "Establish monitoring and tracking procedures"
        ]
        for step in next_steps:
            p = doc.add_paragraph(style='List Bullet')
            p.text = step
        
        # Save document
        doc.save(filepath)
        
        return filepath
    
    def _generate_html_proposal(self, template_name: str) -> str:
        """Generate HTML proposal"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        client_name = self.roi_results.get('inputs', {}).get('company_name', 'Client')
        safe_client = ''.join(c for c in client_name if c.isalnum() or c in (' ', '-', '_')).strip()
        
        filename = f'ROI_Proposal_{safe_client}_{timestamp}.html'
        filepath = os.path.join(self.temp_dir, filename)
        
        # Load template
        template_path = self._get_html_template_path(template_name)
        template_content = self._load_html_template(template_path)
        
        # Generate content
        executive_summary = self._generate_executive_summary()
        implementation_timeline = self._generate_implementation_timeline()
        cost_breakdown = self._generate_cost_breakdown()
        
        # Replace template variables
        html_content = template_content.format(
            client_name=client_name,
            company_name=self.company_config['name'],
            date=datetime.now().strftime('%B %d, %Y'),
            executive_summary=executive_summary.replace('\n\n', '</p><p>'),
            primary_color=self.company_config['primary_color'],
            secondary_color=self.company_config['secondary_color'],
            accent_color=self.company_config['accent_color'],
            # ROI metrics
            initial_investment=f"${self.roi_results.get('inputs', {}).get('service_investment', 0):,.2f}",
            first_year_roi=f"{self.roi_results.get('roi_metrics', {}).get('first_year_roi', 0)*100:.1f}%",
            annual_savings=f"${self.roi_results.get('roi_metrics', {}).get('annual_savings', 0):,.2f}",
            payback_period=f"{self.roi_results.get('roi_metrics', {}).get('payback_period_months', 0):.1f} months",
            npv_3year=f"${self.roi_results.get('financial_metrics', {}).get('npv', 0):,.2f}",
            # Charts (will be embedded as base64)
            chart_roi_timeline=self._get_chart_base64('roi_timeline.png'),
            chart_cost_breakdown=self._get_chart_base64('cost_breakdown.png'),
            chart_savings_projection=self._get_chart_base64('savings_projection.png'),
            # Timeline
            timeline_items=''.join(f'<li>{item}</li>' for item in implementation_timeline),
            # Cost breakdown
            cost_breakdown_html=self._format_cost_breakdown_html(cost_breakdown)
        )
        
        # Write HTML file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filepath
    
    def _get_html_template_path(self, template_name: str) -> str:
        """Get path to HTML template"""
        templates_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
        return os.path.join(templates_dir, f'proposal_template_{template_name}.html')
    
    def _load_html_template(self, template_path: str) -> str:
        """Load HTML template, create default if doesn't exist"""
        if os.path.exists(template_path):
            with open(template_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            # Return default template
            return self._get_default_html_template()
    
    def _get_default_html_template(self) -> str:
        """Get default HTML template"""
        return '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>ROI Analysis Proposal - {client_name}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 0;
            color: #333;
            background-color: #f8f9fa;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 40px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid {primary_color};
            padding-bottom: 30px;
            margin-bottom: 40px;
        }}
        .title {{
            color: {primary_color};
            font-size: 2.5em;
            margin-bottom: 20px;
            font-weight: 300;
            letter-spacing: 2px;
        }}
        .subtitle {{
            font-size: 1.1em;
            color: #666;
            margin: 5px 0;
        }}
        .section {{
            margin-bottom: 40px;
        }}
        .section h2 {{
            color: {primary_color};
            border-left: 4px solid {accent_color};
            padding-left: 20px;
            font-size: 1.8em;
            margin-bottom: 20px;
        }}
        .metrics-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            margin: 30px 0;
        }}
        .metric-card {{
            background: linear-gradient(135deg, {secondary_color}, {primary_color});
            color: white;
            padding: 25px;
            border-radius: 10px;
            text-align: center;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .metric-value {{
            font-size: 2em;
            font-weight: bold;
            margin-bottom: 10px;
        }}
        .metric-label {{
            font-size: 0.9em;
            opacity: 0.9;
        }}
        .chart-container {{
            text-align: center;
            margin: 30px 0;
        }}
        .chart-container img {{
            max-width: 100%;
            height: auto;
            border-radius: 8px;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .timeline {{
            list-style: none;
            padding: 0;
        }}
        .timeline li {{
            background: #f8f9fa;
            margin: 10px 0;
            padding: 15px 20px;
            border-left: 4px solid {accent_color};
            border-radius: 0 8px 8px 0;
        }}
        .cost-breakdown {{
            background: #f8f9fa;
            padding: 25px;
            border-radius: 10px;
            margin: 20px 0;
        }}
        .footer {{
            text-align: center;
            margin-top: 60px;
            padding-top: 30px;
            border-top: 1px solid #eee;
            color: #666;
        }}
        @media print {{
            body {{ background: white; }}
            .container {{ box-shadow: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1 class="title">ROI ANALYSIS PROPOSAL</h1>
            <div class="subtitle">Prepared for: <strong>{client_name}</strong></div>
            <div class="subtitle">Prepared by: <strong>{company_name}</strong></div>
            <div class="subtitle">Date: <strong>{date}</strong></div>
        </div>

        <div class="section">
            <h2>Executive Summary</h2>
            <p>{executive_summary}</p>
        </div>

        <div class="section">
            <h2>Key ROI Metrics</h2>
            <div class="metrics-grid">
                <div class="metric-card">
                    <div class="metric-value">{initial_investment}</div>
                    <div class="metric-label">Initial Investment</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{first_year_roi}</div>
                    <div class="metric-label">First Year ROI</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{annual_savings}</div>
                    <div class="metric-label">Annual Savings</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{payback_period}</div>
                    <div class="metric-label">Payback Period</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{npv_3year}</div>
                    <div class="metric-label">3-Year NPV</div>
                </div>
            </div>
        </div>

        <div class="section">
            <h2>Visual Analysis</h2>
            <div class="chart-container">
                <h3>ROI Timeline & Cumulative Savings</h3>
                <img src="data:image/png;base64,{chart_roi_timeline}" alt="ROI Timeline">
            </div>
            <div class="chart-container">
                <h3>Annual Cost Breakdown</h3>
                <img src="data:image/png;base64,{chart_cost_breakdown}" alt="Cost Breakdown">
            </div>
            <div class="chart-container">
                <h3>36-Month Savings Projection</h3>
                <img src="data:image/png;base64,{chart_savings_projection}" alt="Savings Projection">
            </div>
        </div>

        <div class="section">
            <h2>Implementation Timeline</h2>
            <ul class="timeline">
                {timeline_items}
            </ul>
        </div>

        <div class="section">
            <h2>Detailed Cost Analysis</h2>
            <div class="cost-breakdown">
                {cost_breakdown_html}
            </div>
        </div>

        <div class="section">
            <h2>Next Steps</h2>
            <ul class="timeline">
                <li>Review and approve this ROI analysis</li>
                <li>Schedule implementation kickoff meeting</li>
                <li>Finalize project timeline and milestones</li>
                <li>Begin service deployment and integration</li>
                <li>Establish monitoring and tracking procedures</li>
            </ul>
        </div>

        <div class="footer">
            <p><strong>{company_name}</strong></p>
            <p>Professional ROI Analysis & Business Solutions</p>
        </div>
    </div>
</body>
</html>'''
    
    def _get_chart_base64(self, chart_filename: str) -> str:
        """Convert chart to base64 for HTML embedding"""
        import base64
        
        chart_path = os.path.join(self.temp_dir, chart_filename)
        if os.path.exists(chart_path):
            with open(chart_path, 'rb') as f:
                return base64.b64encode(f.read()).decode('utf-8')
        return ''
    
    def _format_cost_breakdown_html(self, cost_breakdown: Dict[str, str]) -> str:
        """Format cost breakdown for HTML"""
        html = ''
        for section, content in cost_breakdown.items():
            html += f'<h3>{section}</h3><p>{content}</p>'
        return html
    
    def _generate_executive_summary(self) -> str:
        """Generate executive summary text"""
        client_name = self.roi_results.get('inputs', {}).get('company_name', 'your organization')
        investment = self.roi_results.get('inputs', {}).get('service_investment', 50000)
        roi = self.roi_results.get('roi_metrics', {}).get('first_year_roi', 0.25) * 100
        annual_savings = self.roi_results.get('roi_metrics', {}).get('annual_savings', 60000)
        payback_months = self.roi_results.get('roi_metrics', {}).get('payback_period_months', 18)
        
        return f"""This comprehensive ROI analysis demonstrates the significant financial benefits of implementing our business optimization solution for {client_name}.

Our analysis reveals that with an initial investment of ${investment:,.2f}, {client_name} can expect to achieve a first-year ROI of {roi:.1f}%, with annual cost savings of ${annual_savings:,.2f}. The investment will pay for itself in just {payback_months:.1f} months.

The key drivers of these savings include operational efficiency improvements, error reduction, streamlined processes, and enhanced productivity. Our solution addresses current pain points while positioning {client_name} for sustainable long-term growth.

This proposal outlines the detailed financial projections, implementation timeline, and expected outcomes to help {client_name} make an informed decision about this strategic investment."""
    
    def _generate_implementation_timeline(self) -> List[str]:
        """Generate implementation timeline"""
        return [
            "Week 1-2: Project kickoff and initial assessment",
            "Week 3-4: System configuration and customization",
            "Week 5-6: Integration with existing systems",
            "Week 7-8: User training and change management",
            "Week 9-10: Testing and quality assurance",
            "Week 11-12: Go-live and initial monitoring",
            "Month 4-6: Optimization and fine-tuning",
            "Month 7-12: Full deployment and ROI tracking"
        ]
    
    def _generate_cost_breakdown(self) -> Dict[str, str]:
        """Generate detailed cost breakdown"""
        inputs = self.roi_results.get('inputs', {})
        
        labor_annual = inputs.get('labor_costs', 0) * 12
        shipping_annual = inputs.get('shipping_costs', 0) * 12
        error_annual = inputs.get('error_costs', 0) * 12
        inventory_annual = inputs.get('inventory_costs', 0) * 12
        
        return {
            "Current Annual Costs": f"Your organization currently incurs approximately ${labor_annual + shipping_annual + error_annual + inventory_annual:,.2f} annually in operational costs that our solution will optimize.",
            
            "Labor Costs": f"Annual labor costs of ${labor_annual:,.2f} include manual processing, error correction, and inefficient workflows. Our solution can reduce these costs by 25-40% through automation and process optimization.",
            
            "Shipping & Logistics": f"Current shipping costs of ${shipping_annual:,.2f} annually can be optimized through better routing, carrier selection, and bulk shipping strategies, typically yielding 15-25% savings.",
            
            "Error & Return Costs": f"Error-related costs of ${error_annual:,.2f} annually represent a significant optimization opportunity. Our solution typically reduces error rates by 60-80%, directly impacting the bottom line.",
            
            "Inventory Management": f"Inventory carrying costs of ${inventory_annual:,.2f} annually can be reduced through better demand forecasting and inventory optimization, typically achieving 20-30% improvements.",
            
            "Investment Breakdown": f"The total investment of ${inputs.get('service_investment', 0):,.2f} covers software licensing, implementation services, training, and first-year support. This investment is fully justified by the projected savings and ROI."
        }
    
    def cleanup(self) -> None:
        """Clean up temporary files and directories"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temporary directory: {self.temp_dir}")
        except Exception as e:
            logger.warning(f"Could not clean up temp directory {self.temp_dir}: {e}")
    
    def __del__(self):
        """Destructor to ensure cleanup"""
        self.cleanup()


def create_proposal_template(template_name: str = 'professional') -> str:
    """Create HTML proposal template file"""
    templates_dir = os.path.join(os.path.dirname(__file__), '..', 'templates')
    os.makedirs(templates_dir, exist_ok=True)
    
    template_path = os.path.join(templates_dir, f'proposal_template_{template_name}.html')
    
    # Use the default template from ProposalGenerator
    generator = ProposalGenerator({})  # Empty results for template creation
    template_content = generator._get_default_html_template()
    
    with open(template_path, 'w', encoding='utf-8') as f:
        f.write(template_content)
    
    logger.info(f"Created proposal template: {template_path}")
    return template_path


if __name__ == '__main__':
    # Example usage
    sample_results = {
        'inputs': {
            'company_name': 'Acme E-commerce Ltd.',
            'annual_revenue': 2000000,
            'service_investment': 75000,
            'labor_costs': 8000,
            'shipping_costs': 5000,
            'error_costs': 2000,
            'inventory_costs': 3000
        },
        'roi_metrics': {
            'first_year_roi': 0.45,
            'annual_savings': 85000,
            'monthly_savings': 7083,
            'payback_period_months': 10.6
        },
        'financial_metrics': {
            'npv': 185000,
            'irr': 0.52
        },
        'projections': {
            'year_1': {'roi_percentage': 0.45, 'cumulative_savings': 85000},
            'year_2': {'roi_percentage': 0.52, 'cumulative_savings': 178500},
            'year_3': {'roi_percentage': 0.58, 'cumulative_savings': 285000}
        }
    }
    
    # Create proposal generator
    generator = ProposalGenerator(sample_results)
    
    # Generate all formats
    results = generator.generate_proposal('all')
    
    print("Generated proposal files:")
    for format_type, filepath in results.items():
        print(f"  {format_type.upper()}: {filepath}")
    
    # Create template
    create_proposal_template('professional')